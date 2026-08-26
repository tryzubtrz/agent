from __future__ import annotations

import base64
import csv
import hashlib
import io
import json
import zipfile
from typing import Any, Literal
from uuid import uuid4

from docx import Document
from fastapi import APIRouter, Depends, HTTPException
from openpyxl import load_workbook
from pydantic import BaseModel, Field
from pypdf import PdfReader

from .api import _audit, _memory_fallback, memory_store
from .auth import Principal, require_owner
from .secret_safety import redact_secrets

router = APIRouter()
_MAX_BYTES = 4_000_000
_MAX_TEXT = 120_000


class DocumentExtractBody(BaseModel):
    filename: str = Field(min_length=1, max_length=300)
    content_base64: str = Field(min_length=1, max_length=6_000_000)
    kind: Literal["auto", "pdf", "docx", "xlsx", "csv", "text", "zip"] = "auto"
    save_to_knowledge: bool = False


def _decode(body: DocumentExtractBody) -> bytes:
    try:
        data = base64.b64decode(body.content_base64, validate=True)
    except Exception as exc:
        raise HTTPException(status_code=422, detail="Invalid base64 document payload") from exc
    if len(data) > _MAX_BYTES:
        raise HTTPException(status_code=413, detail="Document exceeds the 4 MB parser limit")
    return data


def _kind(filename: str, requested: str) -> str:
    if requested != "auto":
        return requested
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return "pdf"
    if lower.endswith(".docx"):
        return "docx"
    if lower.endswith(".xlsx"):
        return "xlsx"
    if lower.endswith(".csv"):
        return "csv"
    if lower.endswith(".zip"):
        return "zip"
    return "text"


def _extract_pdf(data: bytes) -> tuple[str, dict[str, Any]]:
    try:
        reader = PdfReader(io.BytesIO(data))
        pages: list[str] = []
        for index, page in enumerate(reader.pages[:100]):
            pages.append(f"--- page {index + 1} ---\n{page.extract_text() or ''}")
        return "\n\n".join(pages), {"pages": len(reader.pages), "parser": "pypdf-text"}
    except Exception as exc:
        raise HTTPException(status_code=422, detail="PDF could not be parsed. Scanned PDFs require OCR.") from exc


def _extract_docx(data: bytes) -> tuple[str, dict[str, Any]]:
    try:
        document = Document(io.BytesIO(data))
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table_index, table in enumerate(document.tables[:50]):
            parts.append(f"--- table {table_index + 1} ---")
            for row in table.rows[:500]:
                parts.append("\t".join(cell.text for cell in row.cells))
        return "\n".join(parts), {"paragraphs": len(document.paragraphs), "tables": len(document.tables), "parser": "python-docx"}
    except Exception as exc:
        raise HTTPException(status_code=422, detail="DOCX could not be parsed") from exc


def _extract_xlsx(data: bytes) -> tuple[str, dict[str, Any]]:
    try:
        workbook = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        parts: list[str] = []
        cells = 0
        for sheet in workbook.worksheets[:30]:
            parts.append(f"--- sheet: {sheet.title} ---")
            for row in sheet.iter_rows(max_row=2000, max_col=80, values_only=True):
                values = ["" if value is None else str(value) for value in row]
                if any(values):
                    parts.append("\t".join(values))
                    cells += len(values)
                if cells >= 50_000:
                    break
            if cells >= 50_000:
                break
        return "\n".join(parts), {"sheets": len(workbook.sheetnames), "cells_read": cells, "parser": "openpyxl-readonly"}
    except Exception as exc:
        raise HTTPException(status_code=422, detail="XLSX could not be parsed") from exc


def _extract_csv(data: bytes) -> tuple[str, dict[str, Any]]:
    text = data.decode("utf-8-sig", errors="replace")
    rows: list[str] = []
    for index, row in enumerate(csv.reader(io.StringIO(text))):
        if index >= 5000:
            break
        rows.append("\t".join(row[:100]))
    return "\n".join(rows), {"rows_read": len(rows), "parser": "csv"}


def _inspect_zip(data: bytes) -> tuple[str, dict[str, Any]]:
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            entries: list[dict[str, Any]] = []
            unsafe = False
            expanded = 0
            for info in archive.infolist()[:1000]:
                normalized = info.filename.replace("\\", "/")
                path_parts = [part for part in normalized.split("/") if part]
                suspicious = normalized.startswith("/") or any(part == ".." for part in path_parts)
                unsafe = unsafe or suspicious
                expanded += info.file_size
                entries.append({
                    "name": info.filename,
                    "compressed_bytes": info.compress_size,
                    "uncompressed_bytes": info.file_size,
                    "directory": info.is_dir(),
                    "unsafe_path": suspicious,
                })
            summary = json.dumps(entries, ensure_ascii=False, indent=2)
            return summary, {
                "entries": len(archive.infolist()),
                "listed": len(entries),
                "unsafe_paths": unsafe,
                "uncompressed_bytes": expanded,
                "executed": False,
                "extracted": False,
                "parser": "zip-inspection-only",
            }
    except zipfile.BadZipFile as exc:
        raise HTTPException(status_code=422, detail="ZIP archive is invalid") from exc


def _store(principal: Principal, key: str, value: dict[str, Any]) -> dict[str, Any]:
    if memory_store is not None:
        return memory_store.upsert(
            workspace_id=principal.workspace_id,
            user_id=principal.user_id,
            namespace="documents",
            key=key,
            value=value,
        )
    _memory_fallback[(principal.workspace_id, principal.user_id, "documents", key)] = value
    return {"namespace": "documents", "key": key, "value": value}


@router.post("/api/documents/extract")
def extract_document(body: DocumentExtractBody, principal: Principal = Depends(require_owner)) -> dict[str, Any]:
    data = _decode(body)
    kind = _kind(body.filename, body.kind)
    if kind == "pdf":
        text, metadata = _extract_pdf(data)
    elif kind == "docx":
        text, metadata = _extract_docx(data)
    elif kind == "xlsx":
        text, metadata = _extract_xlsx(data)
    elif kind == "csv":
        text, metadata = _extract_csv(data)
    elif kind == "zip":
        text, metadata = _inspect_zip(data)
    else:
        text = data.decode("utf-8", errors="replace")
        metadata = {"parser": "utf8-text"}

    safe_text, redacted = redact_secrets(text[:_MAX_TEXT])
    digest = hashlib.sha256(data).hexdigest()
    result: dict[str, Any] = {
        "filename": body.filename,
        "kind": kind,
        "sha256": digest,
        "bytes": len(data),
        "text": safe_text,
        "truncated": len(text) > _MAX_TEXT,
        "redacted": redacted,
        "metadata": metadata,
        "saved": False,
    }

    if body.save_to_knowledge:
        if kind == "zip":
            raise HTTPException(status_code=422, detail="ZIP inspection reports are not stored as knowledge automatically")
        document_id = str(uuid4())
        _store(
            principal,
            document_id,
            {
                "id": document_id,
                "filename": body.filename,
                "kind": kind,
                "sha256": digest,
                "text": safe_text,
                "metadata": metadata,
                "redacted": redacted,
            },
        )
        result["saved"] = True
        result["document_id"] = document_id

    _audit(
        principal,
        event_type="document.parsed",
        payload={
            "filename": body.filename[:180],
            "kind": kind,
            "bytes": len(data),
            "redacted": redacted,
            "saved": result["saved"],
        },
    )
    return result


@router.get("/api/documents")
def list_documents(principal: Principal = Depends(require_owner)) -> list[dict[str, Any]]:
    if memory_store is not None:
        return memory_store.list_for_user(
            workspace_id=principal.workspace_id,
            user_id=principal.user_id,
            namespace="documents",
            limit=250,
        )
    return [
        {"namespace": namespace, "key": key, "value": value}
        for (workspace_id, user_id, namespace, key), value in _memory_fallback.items()
        if workspace_id == principal.workspace_id and user_id == principal.user_id and namespace == "documents"
    ][:250]
